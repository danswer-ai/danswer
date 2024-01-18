import io
import tempfile
import os
from datetime import datetime
from datetime import timezone
from typing import Any

import docx
import openpyxl

import msal

from office365.graph_client import GraphClient
from office365.onedrive.sites.site import Site
from office365.onedrive.driveitems.driveItem import DriveItem

from danswer.configs.app_configs import INDEX_BATCH_SIZE
from danswer.configs.constants import DocumentSource
from danswer.connectors.interfaces import GenerateDocumentsOutput
from danswer.connectors.interfaces import LoadConnector
from danswer.connectors.interfaces import PollConnector
from danswer.connectors.interfaces import SecondsSinceUnixEpoch
from danswer.connectors.models import ConnectorMissingCredentialError
from danswer.connectors.models import Document
from danswer.connectors.models import BasicExpertInfo
from danswer.connectors.models import Section
from danswer.utils.logger import setup_logger
from danswer.connectors.cross_connector_utils.file_utils import read_pdf_file

UNSUPPORTED_FILE_TYPE_CONTENT = ""  # idea copied from the google drive side of things


logger = setup_logger()

def get_text_from_xlsx_driveitem(driveitem_object: DriveItem):

    file_content = driveitem_object.get_content().execute_query().value
    excel_file = io.BytesIO(file_content)
    workbook = openpyxl.load_workbook(excel_file, read_only=True)

    full_text = []
    for sheet in workbook.worksheets:
        sheet_string = '\n'.join(
                            ','.join(map(str, row)) 
                            for row in sheet.iter_rows(min_row=1, values_only=True)
                        )
        full_text.append(sheet_string)
        
    return '\n'.join(full_text)

def get_text_from_docx_driveitem(driveitem_object: DriveItem):
    file_content = driveitem_object.get_content().execute_query().value
    full_text = []

    with tempfile.TemporaryDirectory() as local_path:
        with open(os.path.join(local_path, driveitem_object.name), "wb") as local_file:
            local_file.write(file_content)
            doc = docx.Document(local_file.name)
            for para in doc.paragraphs:
                full_text.append(para.text)
    return '\n'.join(full_text)

def get_text_from_pdf_driveitem(driveitem_object: DriveItem):

    file_content = driveitem_object.get_content().execute_query().value
    file_text = read_pdf_file(file=io.BytesIO(file_content), file_name=driveitem_object.name)
    return file_text

def get_text_from_txt_driveitem(driveitem_object: DriveItem):
    file_content: bytes = driveitem_object.get_content().execute_query().value
    text_string = file_content.decode('utf-8')
    return text_string


class SharepointConnector(LoadConnector, PollConnector):
    def __init__(
        self,
        batch_size: int = INDEX_BATCH_SIZE,
        site_list: list[str] = [],
    ) -> None:
        self.batch_size = batch_size
        self.graph_client: GraphClient | None = None
        self.requested_site_list: list[str] = site_list

    def load_credentials(self, credentials: dict[str, Any]) -> dict[str, Any] | None:
        aad_client_id = credentials["aad_client_id"]
        aad_client_secret = credentials["aad_client_secret"]
        aad_directory_id = credentials["aad_directory_id"]

        def _acquire_token_func():
            """
            Acquire token via MSAL
            """
            authority_url = f'https://login.microsoftonline.com/{aad_directory_id}'
            app = msal.ConfidentialClientApplication(
                authority=authority_url,
                client_id=aad_client_id,
                client_credential=aad_client_secret
            )
            token = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
            return token
        
        self.graph_client = GraphClient(_acquire_token_func)
        return None
    
    def _prune_driveitem_list_by_time(
            self, start:datetime, end:datetime, all_driveitem_objects:list[DriveItem]
        ) -> list[DriveItem]:
        
        pruned_driveitem_object_list:list[DriveItem] = []
        for driveitem_object in all_driveitem_objects:
            if driveitem_object.last_modified_datetime>start and driveitem_object.last_modified_datetime<end:
                pruned_driveitem_object_list.append(driveitem_object)
        
        return pruned_driveitem_object_list
    
    def get_all_driveitem_objects(self, site_object_list: list[Site]) -> list[DriveItem]:
        driveitem_list = []
        for site_object in site_object_list:
            site_list_objects = site_object.lists.get().execute_query()
            for site_list_object in site_list_objects:
                try:
                    driveitems = site_list_object.drive.root.get_files(True).execute_query()
                    driveitem_list.extend(driveitems)
                except Exception:
                    pass
                
        return driveitem_list

    def get_all_site_objects(self) -> list[Site]:
        site_object_list: list[Site] = []
        
        sites_object = self.graph_client.sites.get().execute_query()

        if len(self.requested_site_list)>0:
            for requested_site in self.requested_site_list:
                adjusted_string = "/" + requested_site
                for site_object in sites_object:
                    if site_object.web_url.endswith(adjusted_string):
                        site_object_list.append(site_object)
        else:
            site_object_list.extend(sites_object)
        
        return site_object_list
    
    def _fetch_from_sharepoint(
            self, start: datetime | None = None, end: datetime | None = None
    ) -> GenerateDocumentsOutput:
        if self.graph_client is None:
            raise ConnectorMissingCredentialError("Sharepoint")
        
        site_object_list = self.get_all_site_objects()

        driveitem_list = self.get_all_driveitem_objects(site_object_list)

        if start is not None and end is not None:
            driveitem_list = self._prune_driveitem_list_by_time(start, end, driveitem_list)

        # goes over all urls, converts them into Document objects and then yeilds them in batches
        doc_batch: list[Document] = []
        batch_count = 0
        for driveitem_object in driveitem_list:

            doc_batch.append(self.convert_driveitem_object_to_document(driveitem_object))

            batch_count+=1
            if batch_count>=self.batch_size:
                yield doc_batch
                batch_count = 0
                doc_batch: list[Document] = []
        yield doc_batch
    
    def convert_driveitem_object_to_document(
            self, driveitem_object:DriveItem, 
    ) -> Document:
        file_text = self.extract_driveitem_text(driveitem_object)
        doc = Document(
            id=driveitem_object.id, 
            sections=[Section(link=driveitem_object.web_url, text=file_text)],
            source=DocumentSource.SHAREPOINT,
            semantic_identifier=driveitem_object.name,
            doc_updated_at=driveitem_object.last_modified_datetime.replace(tzinfo=timezone.utc), 
            primary_owners=[BasicExpertInfo(display_name=driveitem_object.last_modified_by.user.displayName,
                                            email=driveitem_object.last_modified_by.user.email)],
            metadata={},
        )
        return doc

    def extract_driveitem_text(self, driveitem_object:DriveItem) -> str:
        driveitem_name =  driveitem_object.name
        driveitem_text = UNSUPPORTED_FILE_TYPE_CONTENT

        if driveitem_name.endswith('.docx') :
            driveitem_text = get_text_from_docx_driveitem(driveitem_object)
        elif driveitem_name.endswith('.pdf') :
            driveitem_text = get_text_from_pdf_driveitem(driveitem_object)
        elif driveitem_name.endswith('.xlsx') :
            driveitem_text = get_text_from_xlsx_driveitem(driveitem_object)
        elif driveitem_name.endswith('.txt') :
            driveitem_text = get_text_from_txt_driveitem(driveitem_object)
            
        return driveitem_text
        
    def load_from_state(self) -> GenerateDocumentsOutput:
        return self._fetch_from_sharepoint()

    def poll_source(
        self, start: SecondsSinceUnixEpoch, end: SecondsSinceUnixEpoch
    ) -> GenerateDocumentsOutput:
        start_datetime = datetime.utcfromtimestamp(start)
        end_datetime = datetime.utcfromtimestamp(end)
        return self._fetch_from_sharepoint(start=start_datetime, end=end_datetime)


if __name__ == "__main__":
    import os

    connector = SharepointConnector(
        sites=os.environ["SITES"]
    )

    connector.load_credentials(
        {
            "aad_client_id": os.environ["AAD_CLIENT_ID"],
            "aad_client_secret": os.environ["AAD_CLIENT_SECRET"],
            "aad_directory_id": os.environ["AAD_CLIENT_DIRECTORY_ID"],
        }
    )
    document_batches = connector.load_from_state()
    print(next(document_batches))
